"""
Paper Export Service - Generate exports in PDF, Word, and LaTeX formats
"""
import io
import logging
from typing import Optional
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY

from app.models.paper import Paper, PaperSection

logger = logging.getLogger(__name__)


class PaperExportService:
    """Service for exporting papers in various formats"""

    def export_to_word(self, paper: Paper) -> bytes:
        """
        Export paper to Microsoft Word (.docx) format

        Args:
            paper: Paper model with sections loaded

        Returns:
            bytes: Word document as bytes
        """
        logger.info(f"📄 Exporting paper '{paper.title}' to Word format")

        # Create document
        doc = Document()

        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Add title
        title = doc.add_heading(paper.title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add metadata
        if paper.co_authors:
            authors_text = ", ".join(paper.co_authors)
            authors = doc.add_paragraph(authors_text)
            authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
            authors_run = authors.runs[0]
            authors_run.font.size = Pt(12)

        # Add date
        date_text = datetime.now().strftime("%B %d, %Y")
        date_para = doc.add_paragraph(date_text)
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.runs[0]
        date_run.font.size = Pt(11)
        date_run.font.italic = True

        doc.add_paragraph()  # Spacing

        # Add abstract if exists
        if paper.abstract and paper.abstract.strip():
            doc.add_heading('Abstract', level=1)
            abstract = doc.add_paragraph(paper.abstract)
            abstract.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            doc.add_paragraph()  # Spacing

        # Add sections in order
        sorted_sections = sorted(paper.sections, key=lambda s: s.order)

        for section in sorted_sections:
            # Add section title
            doc.add_heading(section.title, level=1)

            # Add section content
            if section.content and section.content.strip():
                content_para = doc.add_paragraph(section.content)
                content_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            doc.add_paragraph()  # Spacing between sections

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        logger.info(f"✅ Word export completed for paper '{paper.title}'")
        return buffer.getvalue()

    def export_to_pdf(self, paper: Paper) -> bytes:
        """
        Export paper to PDF format

        Args:
            paper: Paper model with sections loaded

        Returns:
            bytes: PDF document as bytes
        """
        logger.info(f"📄 Exporting paper '{paper.title}' to PDF format")

        buffer = io.BytesIO()

        # Create PDF document
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18,
        )

        # Container for the 'Flowable' objects
        elements = []

        # Define styles
        styles = getSampleStyleSheet()

        # Title style
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor='black',
            spaceAfter=30,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )

        # Author style
        author_style = ParagraphStyle(
            'Author',
            parent=styles['Normal'],
            fontSize=12,
            alignment=TA_CENTER,
            spaceAfter=12,
        )

        # Section heading style
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor='black',
            spaceAfter=12,
            spaceBefore=12,
            fontName='Helvetica-Bold'
        )

        # Body text style
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['BodyText'],
            fontSize=11,
            alignment=TA_JUSTIFY,
            spaceAfter=12,
            leading=14,
        )

        # Add title
        elements.append(Paragraph(paper.title, title_style))
        elements.append(Spacer(1, 12))

        # Add authors if exists
        if paper.co_authors:
            authors_text = ", ".join(paper.co_authors)
            elements.append(Paragraph(authors_text, author_style))

        # Add date
        date_text = datetime.now().strftime("%B %d, %Y")
        elements.append(Paragraph(date_text, author_style))
        elements.append(Spacer(1, 24))

        # Add abstract if exists
        if paper.abstract and paper.abstract.strip():
            elements.append(Paragraph("<b>Abstract</b>", heading_style))
            elements.append(Paragraph(paper.abstract, body_style))
            elements.append(Spacer(1, 12))

        # Add sections in order
        sorted_sections = sorted(paper.sections, key=lambda s: s.order)

        for section in sorted_sections:
            # Add section title
            elements.append(Paragraph(section.title, heading_style))

            # Add section content
            if section.content and section.content.strip():
                # Split content into paragraphs
                paragraphs = section.content.split('\n\n')
                for para in paragraphs:
                    if para.strip():
                        elements.append(Paragraph(para.strip(), body_style))

            elements.append(Spacer(1, 12))

        # Build PDF
        doc.build(elements)

        buffer.seek(0)
        logger.info(f"✅ PDF export completed for paper '{paper.title}'")
        return buffer.getvalue()

    def export_to_latex(self, paper: Paper) -> bytes:
        """
        Export paper to LaTeX format (.tex)

        Args:
            paper: Paper model with sections loaded

        Returns:
            bytes: LaTeX document as bytes
        """
        logger.info(f"📄 Exporting paper '{paper.title}' to LaTeX format")

        # Start LaTeX document
        latex_lines = [
            "\\documentclass[11pt,a4paper]{article}",
            "\\usepackage[utf8]{inputenc}",
            "\\usepackage[T1]{fontenc}",
            "\\usepackage{amsmath}",
            "\\usepackage{graphicx}",
            "\\usepackage{hyperref}",
            "\\usepackage[margin=1in]{geometry}",
            "",
            f"\\title{{{self._escape_latex(paper.title)}}}",
        ]

        # Add authors if exists
        if paper.co_authors:
            authors = " \\and ".join([self._escape_latex(author) for author in paper.co_authors])
            latex_lines.append(f"\\author{{{authors}}}")
        else:
            latex_lines.append("\\author{}")

        latex_lines.extend([
            f"\\date{{\\today}}",
            "",
            "\\begin{document}",
            "",
            "\\maketitle",
            "",
        ])

        # Add abstract if exists
        if paper.abstract and paper.abstract.strip():
            latex_lines.extend([
                "\\begin{abstract}",
                self._escape_latex(paper.abstract),
                "\\end{abstract}",
                "",
            ])

        # Add sections in order
        sorted_sections = sorted(paper.sections, key=lambda s: s.order)

        for section in sorted_sections:
            # Add section title
            latex_lines.append(f"\\section{{{self._escape_latex(section.title)}}}")
            latex_lines.append("")

            # Add section content
            if section.content and section.content.strip():
                # Escape LaTeX special characters
                escaped_content = self._escape_latex(section.content)
                latex_lines.append(escaped_content)
                latex_lines.append("")

        # End document
        latex_lines.extend([
            "\\end{document}",
        ])

        # Join lines and convert to bytes
        latex_content = "\n".join(latex_lines)

        logger.info(f"✅ LaTeX export completed for paper '{paper.title}'")
        return latex_content.encode('utf-8')

    def _escape_latex(self, text: str) -> str:
        """
        Escape special LaTeX characters

        Args:
            text: Text to escape

        Returns:
            str: Escaped text
        """
        # LaTeX special characters that need escaping
        replacements = {
            '&': r'\&',
            '%': r'\%',
            '$': r'\$',
            '#': r'\#',
            '_': r'\_',
            '{': r'\{',
            '}': r'\}',
            '~': r'\textasciitilde{}',
            '^': r'\textasciicircum{}',
            '\\': r'\textbackslash{}',
        }

        result = text
        for char, replacement in replacements.items():
            result = result.replace(char, replacement)

        return result


# Global instance
paper_export_service = PaperExportService()
